from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "pool-feasibility-user-guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6770"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body(doc, text="", bold_prefix=None):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(paragraph.add_run(bold_prefix), bold=True)
        set_font(paragraph.add_run(text[len(bold_prefix) :]))
    else:
        set_font(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, after=4)
    set_font(paragraph.add_run(text))
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    style_paragraph(paragraph, after=4)
    set_font(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_font(
        paragraph.add_run(text),
        size={1: 16, 2: 13, 3: 12}[level],
        color=BLUE if level < 3 else DARK_BLUE,
        bold=True,
    )
    return paragraph


def add_callout(doc, label, text, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0)
    set_font(paragraph.add_run(label + " "), color=DARK_BLUE, bold=True)
    set_font(paragraph.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(header, after=0)
    set_font(header.add_run("Pool Feasibility NZ  |  Quick Guide"), size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(footer, after=0)
    set_font(footer.add_run("Internal Auckland POC  |  Preliminary screening only"), size=9, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    style_paragraph(title, after=3)
    set_font(title.add_run("Pool Feasibility NZ"), size=26, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=14)
    set_font(subtitle.add_run("Simple guide for staff property screening"), size=14, color=MUTED)

    add_callout(
        doc,
        "Purpose:",
        "Use this tool to make a fast, evidence-led first assessment of residential pool potential on an Auckland property.",
        fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "Important:",
        "This is preliminary desktop screening. It is not consent, engineering, survey, title, utility-location, or construction-safety advice.",
    )

    add_heading(doc, "Quick workflow", 1)
    for step in (
        "Enter the Auckland property address and select Fetch property data.",
        "Confirm the resolved address and parcel identity. If several addresses are returned, choose the exact match before continuing.",
        "Review the aerial map and active official layers. Use the layer checkboxes to show or hide map evidence while investigating the site.",
        "Choose a pool shell size, drag it to the proposed position, and rotate it in 15-degree steps. Custom size allows manual length and width.",
        "Read the placement status and conflict cards. Treat hard conflicts, unknowns, and advisory imagery findings as items for follow-up.",
        "Open Preview PDF report only after the map and pool placement are set. The report uses the current map view and active visible layers.",
    ):
        add_number(doc, step)

    add_heading(doc, "1. Property and evidence", 1)
    add_body(doc, "The first result sections establish what property is being assessed and what mapped information was available.")
    for item in (
        "Address and parcel: check the resolved LINZ address, parcel ID, appellation, and identity status.",
        "Aerial imagery: wait for Imagery verified before relying on the imagery view. If it fails, use Try imagery again and treat the map as incomplete until it recovers.",
        "Official map layers: each checkbox controls a visible layer. The map can show buildings, contours, planning information, flooding, mapped services, and other returned datasets.",
        "Attribution: keep the displayed source attribution with any screenshot or report handoff.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2. Manual pool placement", 1)
    add_body(doc, "Placement is manual. The selected shell is shown on the map and can be moved directly by dragging it.")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    headers = ("Control or overlay", "What it means")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, after=0)
        set_font(paragraph.add_run(text), color=INK, bold=True)
    rows = (
        ("Compact / compact plus / standard", "Select a configured pool shell for the placement check."),
        ("Custom size", "Enter length and width in metres. Values must be between 0.1 m and 30 m."),
        ("Rotate 15°", "Turn the selected shell in 15-degree increments."),
        ("Pool shell", "The selected pool footprint."),
        ("Construction allowance", "Indicative working allowance around the pool."),
        ("Indicative barrier", "Indicative barrier envelope for early screening."),
        ("Access envelope", "Indicative access area for construction planning."),
    )
    for left, right in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (left, right)):
            paragraph = cell.paragraphs[0]
            style_paragraph(paragraph, after=0)
            set_font(paragraph.add_run(text), bold=left in ("Pool shell", "Construction allowance", "Indicative barrier", "Access envelope"))

    add_heading(doc, "3. Understand placement results", 1)
    add_body(doc, "The placement status is an early screening result, not an approval. Use it to decide what needs checking next.")
    for item in (
        "Hard GIS conflict: the construction allowance intersects a mapped exclusion or constraint. Reposition the pool or escalate the constraint for proper review.",
        "Advisory aerial imagery conflict: imagery suggests a possible issue, but imagery alone is not proof. Confirm the condition onsite.",
        "Unknown or unavailable: missing mapped information does not prove that a constraint is absent.",
        "Unconfirmed parcel: placement controls are unavailable until the legal parcel is confirmed for the address.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "4. Create the PDF report", 1)
    for step in (
        "Open the property map and wait until the imagery and pool placement are visible.",
        "Set the map layers and viewport you want included. The report captures the current visible map, including active imagery, parcel, evidence layers, and placement overlays.",
        "Select Preview PDF report and review all three pages.",
        "Select Download PDF. If server PDF generation is unavailable, use Print / save PDF from the preview.",
    ):
        add_number(doc, step)
    add_callout(
        doc,
        "Report check:",
        "Before sending or saving the report, confirm that the address, parcel, map extent, active layers, pool position, and attribution are all visible and correct.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "5. What the result does not confirm", 1)
    add_body(doc, "Always explain that the assessment remains preliminary. It does not confirm:")
    for item in (
        "title interests, easements, covenants, or consent notices",
        "private service locations or depths, BeforeUdig plans, or onsite utility locating",
        "ground conditions, groundwater, retaining walls, drainage, or exact construction access",
        "final pool barrier design, consent requirements, engineering design, or construction safety",
        "that a pool position is approved or safe to build",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Troubleshooting", 1)
    troubleshooting = (
        ("No map", "Check imagery status, select Try imagery again, and confirm that the property request completed."),
        ("Cannot place a pool", "Confirm that the legal parcel is mapped and confirmed. Unconfirmed parcels do not expose placement controls."),
        ("Report button unavailable", "Open the property map and wait for a current map capture before previewing the report."),
        ("Report looks different", "Return to the map, set the desired layers and viewport, wait for the map to finish updating, then preview again."),
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    for cell, text in zip(table.rows[0].cells, ("Problem", "What to do")):
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, after=0)
        set_font(paragraph.add_run(text), color=INK, bold=True)
    for left, right in troubleshooting:
        cells = table.add_row().cells
        for cell, text in zip(cells, (left, right)):
            paragraph = cell.paragraphs[0]
            style_paragraph(paragraph, after=0)
            set_font(paragraph.add_run(text), bold=left == "Problem")

    add_body(doc, "Current boundary: Auckland-only, internal-only, session-scoped proof of concept. Refreshing or closing the browser can lose the result unless the staff user downloads the session assessment.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Pool Feasibility NZ - Simple Staff Guide"
    doc.core_properties.subject = "How to use the Auckland pool feasibility screening tool"
    doc.core_properties.author = "Royal Glass"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
