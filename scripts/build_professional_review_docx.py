from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "pool-professional-review-guide.md"
OUTPUT = ROOT / "professional-review-guide.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "5B6573"
CAUTION = "FFF8E8"
RISK = "FCEBEB"
WHITE = "FFFFFF"


def set_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("GEOMAP POC  |  PROFESSIONAL REVIEW GUIDE")
    set_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("Internal review draft  |  Page ")
    set_font(run, size=9, color=MUTED)
    add_page_field(footer)


def add_rich_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.widow_control = True
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=10, color=NAVY)
        else:
            run = paragraph.add_run(part)
            set_font(run)
    return paragraph


def add_list_item(doc, text, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = add_rich_paragraph(doc, text, style=style)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, name="Consolas", size=9, color=NAVY)


def add_table(doc, rows):
    if not rows:
        return
    columns = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=columns)
    widths = [9360 // columns] * columns
    if columns == 2:
        widths = [2400, 6960]
    elif columns == 3:
        widths = [2200, 3600, 3560]
    elif columns == 4:
        widths = [1800, 2600, 2500, 2460]
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            repeat_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            if row_index == 0:
                shade(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                shade(cell, LIGHT_GRAY)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", value):
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_font(run, size=9.5, bold=True, color=NAVY)
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    set_font(run, name="Consolas", size=8.5, color=NAVY)
                else:
                    run = p.add_run(part)
                    set_font(run, size=9.5, color=NAVY if row_index == 0 else "000000", bold=row_index == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def parse_markdown(doc, text):
    lines = text.replace("\r\n", "\n").split("\n")
    in_code = False
    code_lines = []
    table_rows = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            row = [item.strip() for item in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-+:?", item) for item in row):
                i += 1
                continue
            table_rows.append(row)
            next_line = lines[i + 1] if i + 1 < len(lines) else ""
            if not (next_line.startswith("|") and next_line.endswith("|")):
                add_table(doc, table_rows)
                table_rows = []
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph(style=f"Heading {level}")
            run = p.add_run(heading.group(2))
            set_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
            i += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.*)$", line)
        bullet = re.match(r"^-\s+(.*)$", line)
        if ordered or bullet:
            add_list_item(doc, (ordered or bullet).group(1), ordered=bool(ordered))
            i += 1
            continue
        if line.startswith("> "):
            p = add_rich_paragraph(doc, line[2:])
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor.from_string(MUTED)
            i += 1
            continue
        if line.startswith("**Purpose:**") or line.startswith("**Status:**"):
            p = add_rich_paragraph(doc, line)
            p.paragraph_format.space_after = Pt(7)
            i += 1
            continue
        add_rich_paragraph(doc, line)
        i += 1


def main():
    doc = Document()
    style_document(doc)
    source = SOURCE.read_text(encoding="utf-8")
    source = source.translate(str.maketrans({
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u00b7": "|",
        "\u00a9": "(c)",
    }))
    parse_markdown(doc, source)
    doc.core_properties.title = "Geomap POC - Professional Review Guide"
    doc.core_properties.subject = "How to use, interpret, and troubleshoot the Auckland pool-feasibility POC"
    doc.core_properties.author = "Royal Glass"
    doc.core_properties.comments = "Internal review document"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
